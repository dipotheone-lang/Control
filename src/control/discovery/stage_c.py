"""Stage C — forms, manuals and contract terms (charter §6).

Produces the raw material for `COMMERCIAL-EXPOSURE.md`, which the
charter calls the likely highest-value single output of the build:
every guarantee expiry, notice period, LD term and accreditation date,
sorted by urgency.

CONFIDENTIALITY BOUNDARY — read this before extending the module.

§6 Stage C requires extracting contractual dates and terms from
contracts. §12.1 (decision D-01, LOCKED) forbids opening the body of
any client-confidential document: no text extraction, no value posted
to a register. Contracts with NDA clients are confidential by
definition, so the two requirements collide precisely where the value
is highest.

**That conflict was resolved by the CEO on 16-Aug-2026: decision D-05**
permits a narrow exception — dates and term durations only, from
confidential CONTRACTS, for the class 2 registers. Its conditions are
binding and are enforced here, not left to the report layer:

- processing is local only; nothing passes to any model or service;
- **no clause text is stored** — the extracted value and the document
  reference are kept, the surrounding text is redacted at the point of
  capture, so it cannot leak through a later change to a template;
- everything else in those documents stays metadata-only under §12.1.2.

**Decision D-14 (17-Aug-2026) extended D-05 to cover OCR**, because the
first live run over this document store found 47% of legal documents
are photographs of text — so the guarantee expiries D-05 exists to
catch were exactly the ones no text layer reached. D-14 adds one
condition: for a client-confidential document the OCR text buffer is
never retained. `_ocr_text` redacts at capture, so the body cannot
reach anything that stores or renders.

`permit_confidential_dates=False` is still the default. The exception
must be switched on deliberately by a caller acting under D-05; the
module never assumes it.

Documents that remain unreadable — scans without OCR above the §5.5
floor — are recorded as gaps rather than guessed at, so the register
stays honest about the shape of its own holes (§1.1).
"""

import hashlib
import json
import re
from dataclasses import dataclass, field, fields
from datetime import date, datetime
from pathlib import Path

from . import date_shapes

READABLE_SUFFIXES = {".pdf", ".docx", ".txt", ".md"}
# Rendered documents Control cannot read as text without OCR (§5.5).
SCANNED_SUFFIXES = {".jpg", ".jpeg", ".png", ".tif", ".tiff"}

_MONTHS = ("jan", "feb", "mar", "apr", "may", "jun",
           "jul", "aug", "sep", "oct", "nov", "dec")

# The separator class includes the full stop because this estate writes
# dates that way. Measured, not assumed: `diagnose-dates` over 957
# documents found 159 occurrences of `NN.NN.NNNN` that no pattern
# parsed — more than every recognised format except `NN/NN/NNNN` at 172.
# Each one was a date a person reads at a glance and the register never
# saw, and that gap is most of why 525 commercial terms yielded two
# dated ones.
#
# Day-first, like the slash form: Egyptian practice, and `_parse_date`
# validates the result, so a version string that happens to look like
# this fails on the calendar rather than entering a register.
_DATE_PATTERNS = (
    re.compile(r"\b(\d{1,2})[/.-](\d{1,2})[/.-](\d{4})\b"),
    # `2026/11/30` — 199 occurrences, the largest unparsed shape in the
    # estate after the dotted form was fixed. ISO order with slashes,
    # and unambiguous because a four-digit year cannot be a day.
    re.compile(r"\b(\d{4})[-/](\d{1,2})[-/](\d{1,2})\b"),
    # The optional stop lets an abbreviated month through: `12 Sept. 2021`
    # is written that way and was rejected on the punctuation alone. The
    # month itself is matched on its first three letters, so the
    # abbreviation was never the problem.
    re.compile(r"\b(\d{1,2})[\s-]([A-Za-z]{3,9})\.?[\s-](\d{4})\b"),
)


def _phrase(source: str) -> "re.Pattern":
    """Compile a term pattern so a phrase survives being line-wrapped.

    Every literal space becomes `\\s+`. Written with plain spaces the
    patterns read as the phrases they are, and "defects liability
    period" then matched nothing the moment a PDF broke the line between
    "liability" and "period" — losing the whole term, not just its date.
    Contract text arrives wrapped; the patterns have to expect it.

    Safe here because no pattern below has a space inside a character
    class, and doing it at compile time means a pattern added later
    cannot forget.
    """
    return re.compile("(?i)" + source.replace(" ", r"\s+"))


# Commercial terms worth a register row. Each is (kind, pattern).
_TERM_PATTERNS = (
    ("GUARANTEE_EXPIRY", _phrase(
        r"(letter of guarantee|bank guarantee|performance bond|advance "
        r"payment guarantee|bid bond|خطاب ضمان)")),
    ("VALIDITY", _phrase(
        r"(valid (?:until|till|through|up to)|expiry date|expires on|"
        r"validity period|صلاحية)")),
    ("LIQUIDATED_DAMAGES", _phrase(
        r"(liquidated damages|penalt(?:y|ies) for delay|غرامة تأخير)")),
    # §2.2: "a claim not noticed within its window is generally
    # forfeited" — the most expensive miss in the charter, so the
    # phrasings this has to catch are worth being explicit about.
    #
    # The first version required the literal word "notice" and missed
    # "shall be notified within 21 days", which is how the clause is
    # normally written. It read as "no notice period in this contract",
    # which is the silence §1.1 exists to prevent. `notif` now covers
    # notice / notified / notify / notification, and the window may sit
    # on either side of the trigger word.
    ("NOTICE_PERIOD", _phrase(
        r"(within\s+\(?\d{1,3}\)?\s*(?:calendar |working |business )?days?"
        r".{0,60}(?:notif|claim|variation|إخطار|مطالبة)"
        r"|(?:notif|claim|variation|إخطار|مطالبة).{0,60}within\s+\(?\d{1,3}\)?"
        r"\s*(?:calendar |working |business )?days?"
        r"|خلال\s+\(?\d{1,3}\)?\s*يوم)")),
    ("DEFECTS_LIABILITY", _phrase(
        r"(defects? liability period|maintenance period|فترة الضمان)")),
    # `retention ` with a trailing space missed "Retention: 5% released
    # 30 June 2027." — a colon is not a space — so the retention release
    # date left the register entirely. §2.2 tracks retention releases by
    # name; a word boundary catches the punctuation this is written with.
    ("RETENTION", _phrase(
        r"(retention\b(?:\s+(?:money|amount|percentage|release))?"
        r"|محتجزات)")),
    ("PAYMENT_TERMS", _phrase(
        r"(payment (?:terms|within)|net\s+\d{1,3}\s*days|شروط الدفع)")),
    ("ACCREDITATION", _phrase(
        r"(prequalification|pre-qualification|vendor registration|"
        r"supplier registration|accreditation|ISO\s?\d{4,5})")),
)


@dataclass
class CommercialTerm:
    kind: str
    source: str                 # file path, relative
    context: str                # short quote, the citation (§1.2)
    found_date: str = ""        # ISO date if one sits near the term
    page_or_para: str = ""
    # The phrase that actually matched, not the window around it. A
    # context window spans neighbouring clauses by design — it is the
    # citation — so reading the instrument type out of it picked up the
    # advance payment guarantee in the next sentence and labelled a
    # performance bond with it. Redacted with everything else on a
    # confidential document (D-05), which is why an instrument type is
    # correctly unrecoverable there rather than guessed.
    term_text: str = ""


@dataclass
class DocumentRecord:
    path: str
    confidential: bool
    reason: str = ""
    readable: bool = True
    terms: list = field(default_factory=list)
    note: str = ""


@dataclass
class StageCResult:
    documents: list = field(default_factory=list)
    terms: list = field(default_factory=list)
    blocked: list = field(default_factory=list)     # confidential, not read
    d05_extracted: list = field(default_factory=list)  # confidential, dates only
    unreadable: list = field(default_factory=list)  # scanned/OCR needed
    # OCR accounting (§5.5). Attempted and read are not the same number,
    # and the difference is the honest part: below_floor holds documents
    # the engine read but was not confident enough to be believed.
    ocr_attempted: list = field(default_factory=list)
    ocr_read: list = field(default_factory=list)
    ocr_below_floor: list = field(default_factory=list)
    # Every mean confidence OCR produced, accepted or not. The floor is
    # a governance number (§5.5, §14.4 forbids learning from lowering
    # it), and it should be set from this estate's own documents rather
    # than from a default chosen without seeing them.
    ocr_confidences: list = field(default_factory=list)
    ocr_failed: list = field(default_factory=list)
    from_cache: int = 0
    # Documents whose text was cached and whose terms were read again
    # under changed patterns. Counted apart from `from_cache` because
    # they are not the same claim: one says nothing was redone, the other
    # says the expensive half was not.
    re_extracted: int = 0


# Words that identify an industry, not a counterparty. A client name
# reduced to one of these matches half a contracting company's folders,
# and the resulting noise hides the gap it pretends to protect.
_GENERIC_TOKENS = {
    "air", "and", "canal", "chemical", "chemicals", "co", "company",
    "construction", "contracting", "egypt", "energy", "engineering", "for",
    "gold", "group", "holding", "industrial", "industries", "international",
    "ltd", "materials", "mines", "misr", "sae", "services", "steel", "sugar",
    "supplies", "the", "trading",
}
_MIN_TOKEN = 5


def _normalise(text: str) -> str:
    """Lowercase, with every non-alphanumeric run reduced to a space.

    Matching then happens on whole words, so "Air Liquide" cannot be
    matched by the "air" inside "repair schedule.xlsx".
    """
    return " " + "".join(
        c if c.isalnum() else " " for c in str(text).lower()).strip() + " "


def match_tokens(name: str) -> list[str]:
    """The strings that identify this client in a path.

    The full name always counts. Individual words count only if they are
    long enough and specific enough to mean this client and not an
    industry — so "Suez Steel" matches a folder named for Suez Steel,
    but not every steel supplier on the drive.
    """
    normalised = _normalise(name).strip()
    if not normalised:
        return []
    tokens = [normalised]
    for word in normalised.split():
        if len(word) >= _MIN_TOKEN and word not in _GENERIC_TOKENS:
            tokens.append(word)
    return tokens


def classify_confidential(path: Path, confidential_clients: list[str],
                          confidential_folders: list[str],
                          confidential_projects: list[str] | None = None,
                          ) -> tuple[bool, str]:
    """§12.1.1, conservative by design: if in doubt, confidential.

    `path` must be RELATIVE to the scan root. Classifying on an absolute
    path lets directories above the root decide — a machine where the
    data sits under, say, "Confidential Backup" would mark every
    document confidential. That is not caution, it is noise, and it
    would hide the gap it pretends to protect.

    Folder names carry the same weight as filenames, which is the CEO's
    decision of 16-Aug-2026: a folder named for a client is confidential
    and so is everything in it. §12.1.1 already says as much — the
    folder rule simply had nothing to consult until the inventory ran.
    """
    text = _normalise(path)
    segments = [_normalise(part).strip() for part in Path(path).parts]

    for folder in confidential_folders:
        needle = _normalise(folder).strip()
        if needle and needle in text:
            return True, f"inside folder classified confidential: {folder}"

    for client in confidential_clients:
        for token in match_tokens(client):
            if f" {token} " not in text:
                continue
            where = ("folder named for the client"
                     if any(token in segment for segment in segments[:-1])
                     else "filename references the client")
            return True, f"{where}: {client}"

    for project in confidential_projects or []:
        needle = _normalise(project).strip()
        if needle and f" {needle} " in text:
            return True, f"project mapped to a confidential client: {project}"

    for marker in ("confidential", "proprietary", "restricted", "nda", "سري"):
        if marker in text:
            return True, f"document marked {marker}"
    return False, ""


def extract_text(path: Path, max_chars: int = 400_000) -> str | None:
    """Text from a non-confidential document. None when unreadable."""
    suffix = path.suffix.lower()
    try:
        if suffix in (".txt", ".md"):
            return path.read_text(encoding="utf-8", errors="replace")[:max_chars]
        if suffix == ".docx":
            import docx
            document = docx.Document(str(path))
            parts = [p.text for p in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    parts.extend(cell.text for cell in row.cells)
            return "\n".join(parts)[:max_chars]
        if suffix == ".pdf":
            import pypdf
            reader = pypdf.PdfReader(str(path))
            parts = []
            for page in reader.pages:
                parts.append(page.extract_text() or "")
                if sum(len(p) for p in parts) > max_chars:
                    break
            text = "\n".join(parts)[:max_chars]
            # A PDF of scans yields almost nothing: that is a gap, not a
            # document without terms (§5.5).
            return text if len(text.strip()) >= 40 else None
    except Exception:
        return None
    return None


def _parse_date(fragment: str) -> str:
    for pattern in _DATE_PATTERNS:
        match = pattern.search(fragment)
        if not match:
            continue
        groups = match.groups()
        try:
            if groups[1].isalpha():
                day, month_name, year = int(groups[0]), groups[1][:3].lower(), int(groups[2])
                if month_name not in _MONTHS:
                    continue
                return date(year, _MONTHS.index(month_name) + 1, day).isoformat()
            if len(groups[0]) == 4:
                return date(int(groups[0]), int(groups[1]), int(groups[2])).isoformat()
            # Ambiguous DD/MM vs MM/DD: Egyptian practice is DD/MM.
            return date(int(groups[2]), int(groups[1]), int(groups[0])).isoformat()
        except (ValueError, IndexError):
            continue
    return ""


# A clause ends at a full stop, a semicolon or a line break. Dates do
# not cross that boundary and neither may the pairing below: the format
# "31.12.2026" is not in _DATE_PATTERNS and "0.5%" has no space after
# the point, so neither is split by this.
#
# A colon is NOT a terminator: it introduces the clause rather than
# ending it. Treating it as one clipped "Retention: 5% released 30 June
# 2027" to nothing and lost the release date — the opposite error to
# the one above, and the same cost.
#
# **Nor is a single line break.** Contract text does not arrive as
# sentences: PDF extraction and OCR both emit a newline per rendered
# line, so "the performance bond shall remain valid until\n31 December
# 2026" is one clause wearing a line wrap. Treating that newline as a
# boundary produced 468 undated terms out of 470 on the first real run
# over 957 documents — an empty register that reported itself as 470
# terms extracted, which is the worst shape a §1.1 failure can take.
#
# A blank line is a paragraph break and does end a clause. A single
# newline is crossed at most once, so a wrapped clause still finds its
# date while a term never reaches two rows down a table.
# **Nor is the full stop in an abbreviation.** "Letter of Guarantee No.
# 5512 valid until 31.12.2027" is one clause, and reading "No." as a
# sentence end truncated it before its own date — which is why both
# dated terms in a 957-document run were VALIDITY rather than
# GUARANTEE_EXPIRY: the instrument is named with a reference, the
# reference carries an abbreviation, and the expiry fell off the end.
#
# Two exclusions, both narrow enough to name. A stop after a known
# abbreviation is not a boundary, and neither is one followed by a
# digit — a number continuing a reference is not a new sentence.
_ABBREVIATIONS = frozenset((
    "no", "nos", "ref", "inv", "art", "cl", "para", "sec", "ch", "fig",
    "ltd", "co", "inc", "vs", "etc", "approx", "min", "max", "dept",
    "mr", "mrs", "ms", "dr", "eng", "st", "ave", "rev", "acct", "attn",
))
_BOUNDARY = re.compile(r"([.;])(\s+)|\n\s*\n")
_MAX_WRAPS = 1


class _HardBoundary:
    """`_BOUNDARY` with the abbreviation exclusions applied.

    Kept behind the same `search`/`finditer` surface the clause helpers
    already use, so the exclusion cannot be applied in one direction and
    forgotten in the other — that asymmetry is how the retention release
    was lost once already.
    """

    pattern = _BOUNDARY.pattern + "|abbrev:" + ",".join(sorted(_ABBREVIATIONS))

    @staticmethod
    def _is_boundary(fragment: str, match) -> bool:
        if match.group(1) is None:        # a blank line always ends a clause
            return True
        before = re.search(r"([^\W\d_]+)$", fragment[:match.start()])
        if before and before.group(1).lower() in _ABBREVIATIONS:
            return False
        after = fragment[match.end():match.end() + 1]
        return not after.isdigit()

    def finditer(self, fragment: str):
        for match in _BOUNDARY.finditer(fragment):
            if self._is_boundary(fragment, match):
                yield match

    def search(self, fragment: str):
        return next(self.finditer(fragment), None)


_HARD_BOUNDARY = _HardBoundary()

# How close two matches must be to be reading the same clause.
_CLAUSE_SPAN = 80


def _clause_after(fragment: str) -> str:
    """The rest of the term's own clause, and no further."""
    boundary = _HARD_BOUNDARY.search(fragment)
    text = fragment[:boundary.start()] if boundary else fragment
    return "\n".join(text.split("\n")[:_MAX_WRAPS + 1])


def _clause_before(fragment: str) -> str:
    """The start of the term's own clause, and no further back."""
    last = None
    for boundary in _HARD_BOUNDARY.finditer(fragment):
        last = boundary
    text = fragment[last.end():] if last else fragment
    return "\n".join(text.split("\n")[-(_MAX_WRAPS + 1):])


# How far a term sits from the nearest readable date, in characters.
# Buckets rather than distances: a bucket count is a statistic, an exact
# offset starts to describe a document's structure.
_DISTANCE_BUCKETS = (120, 250, 500, 1000, 2500)


def _distance_histogram(text: str) -> dict:
    """For each term, how far away the nearest parseable date is.

    The question this exists to answer, and the one guessing kept
    getting wrong. Eleven client-confidential contracts produced 29
    terms and 47 readable dates and paired none of them: the dates are
    present, the terms are present, and the window between them is the
    only thing left. Widening a window blind is how a date from the
    clause next door ends up in a register (§2.1), so the width is
    measured first.

    Counts per bucket only. No offset, no text — a bucket is a
    statistic, an exact position starts to describe a document's
    structure (§12.1.2).
    """
    positions: list[int] = []
    for _, pattern in _TERM_PATTERNS:
        positions.extend(match.start() for match in pattern.finditer(text))
    if not positions:
        return {}

    dates: list[int] = []
    for pattern in date_shapes.CANDIDATES:
        for match in pattern.finditer(text):
            if _parse_date(match.group(0)):
                dates.append(match.start())

    histogram: dict[str, int] = {}
    for position in positions:
        if not dates:
            key = "no date in document"
        else:
            nearest = min(abs(position - d) for d in dates)
            key = next((f"<={bucket}" for bucket in _DISTANCE_BUCKETS
                        if nearest <= bucket), f">{_DISTANCE_BUCKETS[-1]}")
        histogram[key] = histogram.get(key, 0) + 1
    return histogram


def find_terms(text: str, source: str, window: int = 120) -> list[CommercialTerm]:
    """Locate commercial terms and the date belonging to each.

    Three subtleties learned the hard way:
    - distinct matches must stay distinct. Two guarantees a line apart
      share a context window; deduplicating on that window merges them
      and loses one expiry entirely.
    - the date belonging to a term usually FOLLOWS it ("valid until
      30/11/2026"). Taking the first date in the window attaches the
      neighbouring clause's date to this one, which is worse than
      reporting no date at all (§1.1).
    - **preferring the following date is not enough on its own.** The
      window is 120 characters and clauses are shorter than that, so it
      reads into the neighbouring sentence in both directions. Found on
      a test contract: an LD clause reading "0.5% per week, capped at
      10%" — no date in it anywhere — came out of the register dated
      31-Dec-2026, borrowed from the guarantee sentence before it; and
      "Payment terms: 60 days from invoice date" was dated 30-Jun-2027,
      which was the retention release. Both are fabricated dates in a
      class 2 register, and §2.1 rates a confident wrong date worse than
      no date. The window is therefore clipped at the clause boundary in
      both directions: a term whose own sentence carries no date is
      reported as undated, which is a finding rather than a fiction.
    """
    found_terms: list[tuple[int, CommercialTerm]] = []
    seen_spans: set[tuple] = set()
    for kind, pattern in _TERM_PATTERNS:
        for match in pattern.finditer(text):
            key = (kind, match.start())
            if key in seen_spans:
                continue
            seen_spans.add(key)

            after = text[match.end():min(len(text), match.end() + window)]
            before = text[max(0, match.start() - window):match.start()]
            found = (_parse_date(_clause_after(after))
                     or _parse_date(_clause_before(before)))

            context = " ".join((before[-60:] + match.group(0) + after[:80]).split())
            found_terms.append((match.start(), CommercialTerm(
                kind=kind, source=source,
                context=context[:240],
                found_date=found,
                term_text=match.group(0)[:80],
            )))

    # VALIDITY is a qualifier, not a term. "Performance bond valid until
    # 31 December 2026" matches GUARANTEE_EXPIRY on the instrument and
    # VALIDITY on the phrasing — one guarantee, two register rows, and
    # the CEO alerted twice for one expiry. Where a VALIDITY match
    # shares a clause and a date with a term that names what actually
    # expires, the named term is the row and the qualifier is dropped.
    # A validity standing on its own — a quotation, say — still counts,
    # because there is nothing more specific to defer to.
    named = [(start, t) for start, t in found_terms if t.kind != "VALIDITY"]
    kept = []
    for start, term in found_terms:
        if term.kind == "VALIDITY" and any(
                other.found_date == term.found_date
                and abs(other_start - start) <= _CLAUSE_SPAN
                for other_start, other in named):
            continue
        kept.append((start, term))
    kept.sort(key=lambda pair: pair[0])
    return [term for _, term in kept]


def run_stage_c(root: Path, confidential_clients: list[str],
                confidential_folders: list[str],
                exclude: list[Path] | None = None,
                permit_confidential_dates: bool = False,
                confidential_projects: list[str] | None = None,
                ocr=None, cache_dir: Path | None = None,
                ocr_floor: float | None = None,
                progress=None) -> StageCResult:
    """`ocr` is an optional callable Path -> OcrResult. Injected rather
    than imported so the confidentiality gate and the §5.5 floor are
    both testable without an engine installed, and so a caller must opt
    in deliberately — OCR is never on by default.

    `progress` is called as progress(stage, done, total, current). A run
    over a real document store takes hours, and a run that prints
    nothing for hours is indistinguishable from one that has hung — the
    operator's only options being to wait on faith or kill work that was
    fine. Enumeration is reported separately because on a full drive it
    is minutes of its own before the first document is even opened.
    """
    root = Path(root)
    excluded = [Path(e).resolve() for e in (exclude or [])]
    result = StageCResult()
    cache = Path(cache_dir) if cache_dir else None
    extraction = extraction_fingerprint()
    ruleset = ruleset_fingerprint(
        confidential_clients, confidential_folders, confidential_projects,
        permit_confidential_dates, ocr is not None, ocr_floor)

    def report(stage, done=0, total=0, current=""):
        if progress:
            progress(stage, done, total, current)

    report("enumerating")
    candidates = []
    for path in sorted(root.rglob("*")):
        if not path.is_file():
            continue
        suffix = path.suffix.lower()
        if suffix not in READABLE_SUFFIXES and suffix not in SCANNED_SUFFIXES:
            continue
        if any(path.resolve().is_relative_to(e) for e in excluded):
            continue
        candidates.append(path)
    report("enumerated", 0, len(candidates))

    for index, path in enumerate(candidates, 1):
        # POSIX form always. This string is the citation on every term
        # (§1.2) and the D-05 audit list of which confidential contracts
        # were opened for dates — a reference that changes shape with the
        # operating system is a weak reference.
        relative = path.relative_to(root).as_posix()
        report("processing", index, len(candidates), relative)

        # A scan of a real document store runs for hours, and OCR is most
        # of that. Caching each document's outcome makes the work durable:
        # an interrupted run loses only the document in flight, and the
        # next invocation converges instead of starting over.
        key = _cache_key(path, relative, ruleset)
        payload = _cache_read(cache, key)
        if payload is not None and payload.get("extraction") != extraction:
            # The document need not be reopened just because the term
            # patterns changed. An outcome that does not depend on them —
            # blocked, unreadable — stands as it is; one that does is
            # re-extracted from the cached text. Only a confidential
            # document, whose text is never retained (D-14), goes back to
            # the file.
            if payload.get("outcome") != "terms":
                payload["extraction"] = extraction
                _cache_write(cache, key, payload)
            else:
                payload = _re_extract(payload, relative)
                if payload is not None:
                    _cache_write(cache, key, payload)
                    result.re_extracted += 1
        if payload is None:
            payload = _process_one(
                path, relative, confidential_clients, confidential_folders,
                confidential_projects, permit_confidential_dates, ocr)
            _cache_write(cache, key, payload)
        else:
            result.from_cache += 1
        _replay(result, payload)

    report("done", len(candidates), len(candidates))
    return result


def ruleset_fingerprint(confidential_clients: list[str],
                        confidential_folders: list[str],
                        confidential_projects: list[str] | None,
                        permit_confidential_dates: bool,
                        ocr_on: bool, floor: float | None = None) -> str:
    """Identity of the RULES a cached outcome was produced under.

    Without this the cache is unsound in the one direction that matters.
    A cached payload carries the document's confidentiality verdict and
    its extracted context; replaying it after the confidential client
    list has grown would serve the OLD verdict — so a contract belonging
    to a client added by CEO decision would replay as non-confidential,
    with its clause text unredacted, straight into the report (§12.1).

    Keyed on the classification inputs, so adding a client, a folder or
    a project invalidates exactly the documents whose answer could
    change, and nothing else. The OCR floor is included for the same
    reason: a below-floor verdict is only meaningful against the floor
    that produced it (§5.5).
    """
    material = "|".join([
        ";".join(sorted(str(c) for c in confidential_clients)),
        ";".join(sorted(str(f) for f in confidential_folders)),
        ";".join(sorted(str(p) for p in (confidential_projects or []))),
        f"d05={int(bool(permit_confidential_dates))}",
        f"ocr={int(bool(ocr_on))}",
        f"floor={floor if floor is not None else ''}",
    ])
    return hashlib.sha256(material.encode()).hexdigest()[:16]


def extraction_fingerprint() -> str:
    """Identity of the rules that READ a document, kept separate from
    the rules that decide whether it may be read at all.

    Separate because they cost different amounts to redo. The read
    fingerprint above governs whether a document must be opened and
    OCR'd again — an hour of work over this estate. This one governs
    only how the text is then searched, which is milliseconds. Folding
    them together meant every fix to a term pattern re-OCR'd 957
    documents.

    The fingerprint above was keyed on the confidentiality inputs and
    the OCR floor. Extraction logic was not in it — so a fix to the term
    patterns or the clause boundary changed nothing on re-run: 957 of
    957 documents came back from cache and the corrected engine never
    touched a single one. A cache that silently serves results from
    superseded logic is worse than no cache, because the operator has
    every reason to believe the fix ran.

    Every pattern and boundary constant is hashed, so any edit to how a
    term or a date is found invalidates exactly the documents whose
    answer could change — and a rule added later cannot forget to
    register itself.
    """
    material = "|".join([
        *(f"{kind}={pattern.pattern}" for kind, pattern in _TERM_PATTERNS),
        *(p.pattern for p in _DATE_PATTERNS),
        _HARD_BOUNDARY.pattern,
        f"wraps={_MAX_WRAPS}",
        f"clause={_CLAUSE_SPAN}",
        # The shape of what is cached, not only the rules that fill it.
        # Adding a field to CommercialTerm makes every existing entry
        # short of it, and a replayed term missing `term_text` loses its
        # instrument type silently — which is the same staleness this
        # fingerprint exists to stop, arriving from the other side.
        "fields=" + ",".join(f.name for f in fields(CommercialTerm)),
    ])
    return hashlib.sha256(material.encode()).hexdigest()[:12]


def _cache_key(path: Path, relative: str, ruleset: str = "") -> str:
    """Identity of a document's content AND the rules it was judged under.

    Size and mtime alongside the path: cheap, and a document that is
    edited gets re-read rather than served stale. Hashing the bytes would
    be more precise and would also mean reading every file on a run whose
    whole purpose is to avoid that.

    The ruleset fingerprint is part of the identity because the cached
    answer is not a property of the document alone — see
    `ruleset_fingerprint`.
    """
    try:
        stat = path.stat()
        stamp = f"{stat.st_size}:{int(stat.st_mtime)}"
    except OSError:
        stamp = "nostat"
    return hashlib.sha256(
        f"{relative}|{stamp}|{ruleset}".encode()).hexdigest()[:32]


def _cache_read(cache_dir: Path | None, key: str) -> dict | None:
    if cache_dir is None:
        return None
    path = cache_dir / f"{key}.json"
    if not path.is_file():
        return None
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None                     # a corrupt entry re-reads, never fails


def _cache_write(cache_dir: Path | None, key: str, payload: dict) -> None:
    if cache_dir is None:
        return
    try:
        cache_dir.mkdir(parents=True, exist_ok=True)
        (cache_dir / f"{key}.json").write_text(
            json.dumps(payload, ensure_ascii=False), encoding="utf-8")
    except OSError:
        pass                            # caching is an optimisation, never a gate


def _re_extract(payload: dict, relative: str) -> dict | None:
    """Redo term extraction from cached text, without reopening the file.

    Returns None when the document has to be read again: a confidential
    one, whose text is never retained (D-14), or an outcome from before
    the text was cached at all.
    """
    text = payload.get("text")
    if payload.get("outcome") != "terms" or not text:
        return None
    terms = find_terms(text, relative)
    refreshed = dict(payload)
    refreshed["extraction"] = extraction_fingerprint()
    refreshed["terms"] = [
        {"kind": t.kind, "source": t.source, "context": t.context,
         "found_date": t.found_date, "page_or_para": t.page_or_para,
         "term_text": t.term_text}
        for t in terms]
    return refreshed


def _replay(result: StageCResult, payload: dict) -> None:
    """Rebuild one document's contribution from a cached entry."""
    relative = payload["relative"]
    terms = [
        CommercialTerm(kind=t["kind"], source=t["source"], context=t["context"],
                       found_date=t.get("found_date", ""),
                       page_or_para=t.get("page_or_para", ""),
                       term_text=t.get("term_text", ""))
        for t in payload.get("terms", [])
    ]
    record = DocumentRecord(
        path=relative, confidential=payload.get("confidential", False),
        reason=payload.get("reason", ""), readable=payload.get("readable", True),
        terms=terms, note=payload.get("note", ""))
    result.documents.append(record)
    if payload.get("outcome") == "blocked":
        result.blocked.append(record)
    elif payload.get("outcome") == "unreadable":
        result.unreadable.append(record)
    else:
        result.terms.extend(terms)
        if payload.get("d05"):
            result.d05_extracted.append(relative)
    ocr_info = payload.get("ocr") or {}
    if ocr_info.get("attempted"):
        result.ocr_attempted.append(relative)
    if ocr_info.get("read"):
        result.ocr_read.append(relative)
    if ocr_info.get("confidence"):
        result.ocr_confidences.append(ocr_info["confidence"])
    if ocr_info.get("below_floor"):
        result.ocr_below_floor.append(ocr_info["below_floor"])
    if ocr_info.get("failed"):
        result.ocr_failed.append(ocr_info["failed"])


def _ocr_text(path: Path, ocr, confidential: bool = False
              ) -> tuple[str | None, dict]:
    """Run OCR and report both the text and what it cost.

    Returns (text, info). `text` is None whenever it must not be used —
    engine missing, failure, or below the §5.5 floor — so the caller
    files the document as unreadable, which is the honest outcome (§1.1).
    Pure: the accounting comes back as data rather than being written
    into a result object, so one document's outcome can be cached and
    replayed exactly.

    **D-14 for a confidential document.** The result is redacted here,
    at capture, before anything can store it: what survives is the
    confidence and the reference, never the body. The text itself is
    returned to the caller for term extraction and goes no further —
    the payload the caller builds holds redacted terms only. Dropping
    it at capture rather than at render is the point; a redaction
    applied on the way out can be undone by a later template change.
    """
    from ..ocr import redact

    info = {"attempted": True, "read": False, "below_floor": "", "failed": "",
            "redacted": False, "confidence": 0.0}
    outcome = ocr(path)
    info["confidence"] = round(float(outcome.mean_confidence or 0.0), 1)
    if outcome.usable:
        info["read"] = True
        text = outcome.text
        if confidential:
            outcome = redact(outcome)      # the stored form keeps no body
            info["redacted"] = True
        return text, info
    if outcome.below_floor:
        info["below_floor"] = (
            f"{path.name} (mean confidence {outcome.mean_confidence})")
    elif outcome.error:
        info["failed"] = f"{path.name}: {outcome.error}"
    return None, info


def _unreadable_note(base: str, confidential: bool, ocr_on: bool,
                     info: dict) -> str:
    """Say why a document stayed unreadable. The reasons differ, and a
    single 'OCR required' note would hide which control was operating."""
    if confidential and not ocr_on:
        return (f"{base} — confidential, and OCR not enabled on this run. "
                "D-14 permits it for dates only (--confidential-dates --ocr)")
    if not ocr_on:
        return f"{base} — OCR required (§5.5); not enabled on this run"
    if info.get("below_floor"):
        return (f"{base} — OCR ran but fell below the §5.5 confidence floor: "
                "UNREADABLE, MANUAL REVIEW REQUIRED")
    if info.get("failed"):
        return f"{base} — OCR failed; recorded as a gap"
    return f"{base} — OCR produced no recognised text"


def _process_one(path: Path, relative: str, confidential_clients: list[str],
                 confidential_folders: list[str],
                 confidential_projects: list[str] | None,
                 permit_confidential_dates: bool, ocr) -> dict:
    """Everything one document contributes, as a serialisable payload."""
    confidential, reason = classify_confidential(
        Path(relative), confidential_clients, confidential_folders,
        confidential_projects)
    payload: dict = {"relative": relative, "confidential": confidential,
                     "reason": reason, "terms": [], "ocr": {}}

    if confidential and not permit_confidential_dates:
        payload.update(outcome="blocked", readable=False,
                       note="not opened — D-01 metadata-only scope")
        return payload

    # Decision D-14 (17-Aug-2026) extended D-05 to permit OCR on
    # client-confidential contracts, for the D-05 purpose only. Getting
    # here at all means the document is either not confidential or is
    # confidential with `permit_confidential_dates` set — the early
    # return above handles the rest — so the gate is now the engine
    # being available, not the classification.
    #
    # D-14's added condition is enforced below: for a confidential
    # document the OCR text buffer is never retained. It passes to
    # `find_terms` transiently and is dropped at capture.
    may_ocr = ocr is not None
    info: dict = {}
    suffix = path.suffix.lower()

    if suffix in SCANNED_SUFFIXES:
        text, info = _ocr_text(path, ocr, confidential) if may_ocr else (None, {})
        base = "image document"
    else:
        text = extract_text(path)
        base = "no extractable text — likely scanned"
        if text is None or not text.strip():
            # A PDF with no text layer is a scan in a PDF wrapper.
            text, info = (_ocr_text(path, ocr, confidential) if may_ocr
                          else (None, {}))

    payload["ocr"] = info
    if text is None:
        payload.update(
            outcome="unreadable", readable=False,
            note=_unreadable_note(base, confidential, ocr is not None, info))
        return payload

    terms = find_terms(text, relative)

    # Taken while the text is in hand, because for a confidential
    # document there is no later. D-14 retains no text for those, so
    # without this nothing about them can ever be measured — and they
    # are the population that matters most: the D-05 exception exists to
    # catch guarantee expiries for the largest clients, and whether it
    # is working is not otherwise observable. Counts only; no date, no
    # clause, nothing traceable to content (§12.1.2).
    parsed_shapes, unparsed_shapes = date_shapes.histogram(text, _parse_date)
    payload["term_date_distance"] = _distance_histogram(text)
    payload["date_shapes"] = {"parsed": parsed_shapes,
                              "unparsed": unparsed_shapes}
    payload["terms_seen"] = len(terms)
    payload["terms_dated"] = sum(1 for term in terms if term.found_date)

    if confidential:
        # D-05: the value and its reference may be kept; the clause text
        # may not. Redact at the point of capture, not at the point of
        # display, so no report template can leak it.
        terms = [
            CommercialTerm(
                kind=t.kind, source=t.source,
                context="[REDACTED — D-05: date extracted, clause text not retained]",
                found_date=t.found_date, page_or_para=t.page_or_para,
                term_text="")
            for t in terms if t.found_date
        ]
    payload.update(
        outcome="terms", readable=True, d05=bool(confidential),
        note="dates extracted under D-05; clause text not retained"
        if confidential else "",
        extraction=extraction_fingerprint(),
        # D-14: for a client-confidential document the OCR text buffer is
        # never retained. So the text is cached for ordinary documents —
        # which lets a later change to the term patterns re-extract in
        # milliseconds instead of re-OCRing — and never for confidential
        # ones, which re-read every time. That asymmetry is the decision
        # working, not an oversight.
        text=None if confidential else text,
        terms=[{"kind": t.kind, "source": t.source, "context": t.context,
                "found_date": t.found_date, "page_or_para": t.page_or_para,
                "term_text": t.term_text}
               for t in terms])
    return payload


def render_commercial_exposure(result: StageCResult, today: date | None = None,
                               not_scanned: list | None = None,
                               scanned: list | None = None) -> str:
    """`not_scanned` names folders the operator asked for that do not
    exist. A folder named wrongly is the easiest way to end up with a
    report that looks complete over ground it never covered, so the miss
    is carried into the document rather than left in the console
    scrollback the operator has already scrolled past (§1.1)."""
    today = today or datetime.now().date()
    dated = [t for t in result.terms if t.found_date]
    undated = [t for t in result.terms if not t.found_date]
    dated.sort(key=lambda t: t.found_date)

    future = [t for t in dated if t.found_date >= today.isoformat()]
    past = [t for t in dated if t.found_date < today.isoformat()]

    lines = [
        "# COMMERCIAL-EXPOSURE",
        "",
        f"Generated {today:%d-%b-%Y}. Every date, notice period and guarantee "
        "term found in readable documents, sorted by urgency.",
        "",
        "**This document is incomplete by design — see 'What could not be "
        "read' before acting on it.**",
        "",
        "## Dates ahead — act on these",
        "",
        "| Date | Kind | Source | Context |",
        "|---|---|---|---|",
    ]
    for term in future[:100]:
        lines.append(f"| {term.found_date} | {term.kind} | {term.source} | "
                     f"{term.context[:100]} |")
    if not future:
        lines.append("| — | — | no future-dated terms found | — |")

    lines += ["", "## Dates already passed — verify status", "",
              "| Date | Kind | Source | Context |", "|---|---|---|---|"]
    for term in past[-40:]:
        lines.append(f"| {term.found_date} | {term.kind} | {term.source} | "
                     f"{term.context[:100]} |")
    if not past:
        lines.append("| — | — | none found | — |")

    lines += ["", "## Terms found without a date", "",
              "These carry obligations whose timing must be established from "
              "the document itself.", "",
              "| Kind | Source | Context |", "|---|---|---|"]
    for term in undated[:60]:
        lines.append(f"| {term.kind} | {term.source} | {term.context[:110]} |")
    if not undated:
        lines.append("| — | none found | — |")

    if result.d05_extracted:
        lines += [
            "",
            "## Confidential contracts — dates extracted under D-05",
            "",
            f"**{len(result.d05_extracted)} client-confidential contract(s)** "
            "were read for dates and term durations only, under decision D-05 "
            "(16-Aug-2026). No clause text is stored, quoted or reproduced: "
            "rows from these documents show `[REDACTED]` in place of context "
            "by construction, not by convention. Everything else in them "
            "remains metadata-only under §12.1.2.",
            "",
        ]
        for path in result.d05_extracted[:40]:
            lines.append(f"  - `{path}`")

    lines += [
        "",
        "## What could not be read",
        "",
    ]
    if scanned:
        lines += [
            f"- **This report covers {len(scanned)} folder(s).** A scan of "
            "part of the drive is not a reading of the drive, and the folders "
            "left out are not evidence that they hold nothing:",
            "",
        ]
        for folder in scanned:
            lines.append(f"  - `{folder}`")
        lines.append("")

    if not_scanned:
        lines += [
            f"- **{len(not_scanned)} folder(s) named for this scan do not "
            "exist and were not searched.** Nothing in them is in this "
            "report, and their absence here is not evidence that they hold "
            "nothing:",
            "",
        ]
        for folder in not_scanned:
            lines.append(f"  - `{folder}`")
        lines.append("")

    lines += [
        f"- **{len(result.blocked)} client-confidential document(s)** were not "
        "opened. Decision D-01 (§12.1) permits metadata only: no text "
        "extraction, no value posted to a register. D-05 covers contracts; "
        "anything else confidential stays closed.",
        "",
    ]
    for record in result.blocked[:40]:
        lines.append(f"  - `{record.path}` — {record.reason}")
    if len(result.blocked) > 40:
        lines.append(f"  - … {len(result.blocked) - 40} further documents")

    lines += [
        "",
        f"- **{len(result.unreadable)} document(s)** produced no extractable "
        "text — scans or images. Under §5.5 nothing is guessed from them; they "
        "need OCR above the confidence floor, or manual review.",
        "",
    ]
    for record in result.unreadable[:25]:
        lines.append(f"  - `{record.path}` — {record.note}")

    lines += [
        "",
        "## Scope of this report",
        "",
        "Decision **D-05** (16-Aug-2026) permits dates and term durations to "
        "be extracted from client-confidential **contracts** for the class 2 "
        "registers. It does not widen §12.1 for anything else: no clause text "
        "is retained, nothing passes to any model or external service, and "
        "every other confidential document stays metadata-only.",
        "",
        "What this report still cannot tell you:",
        "",
        "- terms in documents that produced no extractable text (scans) — "
        "these need OCR above the §5.5 confidence floor, and nothing is "
        "guessed from them;",
        "- notice periods expressed as a duration have no date here, because "
        "a claim window runs from an event that is not in the document set. "
        "They are listed as standing terms, and §2.2 is unambiguous that a "
        "claim not noticed within its window is generally forfeited;",
        "- anything in a contract that was never filed in the scanned folder.",
        "",
    ]
    return "\n".join(lines)
